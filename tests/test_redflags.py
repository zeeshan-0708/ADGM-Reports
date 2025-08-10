# simple test for redflags
from app.redflags import detect_issues_for_doc
from docx import Document
from io import BytesIO

def make_doc(text):
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return Document(bio)

def test_jurisdiction_flag():
    d = make_doc("This agreement is governed by the Dubai Courts.")
    issues = detect_issues_for_doc(d, "Articles of Association")
    assert any("jurisdiction" in i['section'].lower() for i in issues)
