# processor.py
import tempfile
from docx import Document
from io import BytesIO
from typing import List, Dict
from .redflags import detect_issues_for_doc
from .rag import RAGClient
from .checklist import get_required_docs_for_process, infer_process_from_doc_types
from .utils import docx_text, find_clause_locations, get_paragraph_index_by_text
from .comments import add_comment_to_paragraph
import os
import json

# naive keyword detection; improved using RAG/Gemini fallback
def detect_document_type(doc: Document, rag_client: RAGClient=None) -> str:
    text = docx_text(doc).lower()
    if "articles of association" in text or "\naoa" in text or "articles of association" in text:
        return "Articles of Association"
    if "memorandum of association" in text or "moa" in text:
        return "Memorandum of Association"
    if "ubo" in text or "ultimate beneficial owner" in text:
        return "UBO Declaration Form"
    if "register of members" in text or "register of directors" in text:
        return "Register of Members and Directors"
    if "incorporation application" in text or "application for incorporation" in text:
        return "Incorporation Application Form"
    # fallback to RAG/LLM classification if available
    if rag_client:
        q = f"Classify this document into one of: Articles of Association, Memorandum of Association, Incorporation Application Form, UBO Declaration Form, Register of Members and Directors, Other. Document excerpt:\n\n{text[:2000]}"
        resp = rag_client.query_with_context(q, local_context=text[:2000])
        # simple heuristics from reply:
        for dt in ["Articles of Association","Memorandum of Association","Incorporation Application Form","UBO Declaration Form","Register of Members and Directors"]:
            if dt.lower() in resp.lower():
                return dt
    return "Unknown Document Type"

def process_single_docx(file_obj, rag_client: RAGClient=None):
    # read docx into python-docx Document
    file_obj.seek(0)
    doc = Document(file_obj)

    doc_type = detect_document_type(doc, rag_client=rag_client)
    text = docx_text(doc)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    # detect red flag issues
    issues = detect_issues_for_doc(doc, doc_type)

    # Map issues to clause/paragraphs where possible
    clause_map = find_clause_locations(doc)

    # For each issue, try to find a paragraph to attach a comment to:
    annotated_issues = []
    for iss in issues:
        attached = False
        # try to locate keywords from issue text in doc paragraphs
        keywords = []
        # if jurisdiction issue, search for 'jurisdiction' phrase
        if 'jurisdiction' in iss['issue'].lower():
            keywords = ['jurisdiction', 'court', 'adgm']
        elif 'signature' in iss['issue'].lower():
            keywords = ['signature', 'signed by', 'for and on behalf', 'authorized signatory']
        elif 'ubo' in iss['issue'].lower():
            keywords = ['ubo', 'ultimate beneficial owner', 'ownership']
        else:
            # fallback: first paragraph
            keywords = []

        found_para_idx = None
        for kw in keywords:
            idx = get_paragraph_index_by_text(doc, kw, start_from=0)
            if idx is not None:
                found_para_idx = idx
                break

        # If not found, attach to first clause paragraph if any
        if found_para_idx is None and clause_map:
            found_para_idx = clause_map[0][1]

        # If still None, attach to doc start (0)
        if found_para_idx is None:
            found_para_idx = 0 if len(doc.paragraphs) > 0 else None

        # add Word comment using comments.add_comment_to_paragraph
        if found_para_idx is not None:
            target_para = doc.paragraphs[found_para_idx]
            comment_text = f"Issue: {iss['issue']} | Suggestion: {iss.get('suggestion','Refer to ADGM references')}"
            try:
                add_comment_to_paragraph(doc, target_para, comment_text, author="ADGM-Agent", initials="AA")
                attached = True
            except Exception as e:
                # fallback: append text paragraph at end
                doc.add_paragraph(f"[COMMENT] {comment_text}")

        annotated_issues.append({**iss, "attached_to_paragraph_index": found_para_idx, "attached": attached})

    # Use RAG/LMM to contextualize suggestions (limit)
    rag_responses = []
    if rag_client and len(annotated_issues) > 0:
        for iss in annotated_issues[:6]:
            question = f"Document type: {doc_type}. Issue: {iss['issue']}. Provide ADGM-citation and recommended fix."
            rag_answer = rag_client.query_with_context(question, local_context=text[:2000])
            rag_responses.append({"issue": iss, "rag_response": rag_answer})

            # Add a small paragraph with short recommendation next to comment (append)
            doc.add_paragraph(f"RAG Suggestion for issue '{iss['issue']}':")
            doc.add_paragraph(rag_answer[:800])

    # Save the reviewed docx to bytes
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    report = {
        "filename": getattr(file_obj, "name", "uploaded_doc.docx"),
        "document_type": doc_type,
        "issues_found": annotated_issues,
        "num_paragraphs": len(doc.paragraphs)
    }
    return bio.read(), report

def process_uploaded_docs(uploaded_files, use_rag=True, process_choice="Auto-detect"):
    rag_client = RAGClient() if use_rag else None

    reviewed_docs = []
    reports = []
    detected_types = []

    for f in uploaded_files:
        content_bytes, report = process_single_docx(f, rag_client=rag_client)
        reviewed_docs.append({"filename": report["filename"], "content_bytes": content_bytes})
        reports.append(report)
        detected_types.append(report["document_type"])

    # Determine process
    if process_choice == "Auto-detect":
        process_target = infer_process_from_doc_types(detected_types)
    else:
        process_target = process_choice

    required_docs = get_required_docs_for_process(process_target)
    missing = [d for d in required_docs if d not in detected_types]
    summary = {
        "process": process_target,
        "documents_uploaded": len(uploaded_files),
        "required_documents": len(required_docs),
        "missing_documents": missing,
        "individual_reports": reports
    }
    return {"reviewed_docs": reviewed_docs, "summary": summary}
