import re
from docx import Document
from typing import List, Tuple

CLAUSE_REGEX = re.compile(r'(?i)(?:clause\s*)?(\d+(?:\.\d+)+)')

def docx_text(doc: Document) -> str:
    """Return all paragraphs text concatenated."""
    return "\n".join([p.text for p in doc.paragraphs if p.text is not None])

def find_clause_locations(doc: Document) -> List[Tuple[str, int]]:
    """
    Find clause-like headings in doc paragraphs.
    Returns list of (clause_number, paragraph_index)
    """
    results = []
    for idx, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        m = CLAUSE_REGEX.search(p.text)
        if m:
            clause = m.group(1)
            results.append((clause, idx))
    return results

def get_paragraph_index_by_text(doc, text_snippet, start_from=0):
    for i, p in enumerate(doc.paragraphs[start_from:], start=start_from):
        if text_snippet.strip() and text_snippet.lower() in p.text.lower():
            return i
    return None
