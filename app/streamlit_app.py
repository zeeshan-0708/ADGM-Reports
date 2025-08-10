import streamlit as st
from docx.shared import Pt
from docx import Document
import io
import json
import re
import os
from datetime import datetime, timezone
from typing import List, Dict
import time
import pandas as pd

# --- Page Configuration ---
st.set_page_config(
    page_title="ADGM Corporate Agent",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS for Impressive Design ---
st.markdown("""
<style>
/* Import modern fonts */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

/* Global styles */
html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

.main > div {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    min-height: 100vh;
}

/* Custom header styles */
.hero-section {
    background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
    padding: 3rem 2rem;
    border-radius: 20px;
    margin: 1rem 0 2rem 0;
    text-align: center;
    color: white;
    box-shadow: 0 20px 40px rgba(0,0,0,0.1);
}

.hero-title {
    font-size: 3.5rem;
    font-weight: 700;
    margin-bottom: 1rem;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    animation: fadeInUp 1s ease-out;
}

.hero-subtitle {
    font-size: 1.3rem;
    font-weight: 300;
    opacity: 0.9;
    margin-bottom: 2rem;
    animation: fadeInUp 1s ease-out 0.2s both;
}

.hero-stats {
    display: flex;
    justify-content: center;
    gap: 3rem;
    margin-top: 2rem;
    flex-wrap: wrap;
}

.stat-item {
    text-align: center;
    animation: fadeInUp 1s ease-out 0.4s both;
}

.stat-number {
    font-size: 2.5rem;
    font-weight: 700;
    display: block;
    color: #64ffda;
}

.stat-label {
    font-size: 0.9rem;
    opacity: 0.8;
    text-transform: uppercase;
    letter-spacing: 1px;
}

/* Modern card styles */
.custom-card {
    background: rgba(255, 255, 255, 0.95);
    backdrop-filter: blur(10px);
    border-radius: 15px;
    padding: 2rem;
    margin: 1rem 0;
    box-shadow: 0 8px 32px rgba(0,0,0,0.1);
    border: 1px solid rgba(255, 255, 255, 0.2);
    transition: transform 0.3s ease, box-shadow 0.3s ease;
}

.custom-card:hover {
    transform: translateY(-2px);
    box-shadow: 0 15px 45px rgba(0,0,0,0.15);
}

/* Success/Error styles */
.success-item {
    background: linear-gradient(135deg, #10ac84 0%, #1dd1a1 100%);
    color: white;
    padding: 1rem 1.5rem;
    border-radius: 10px;
    margin: 0.5rem 0;
    display: flex;
    align-items: center;
    gap: 1rem;
    animation: slideInLeft 0.5s ease-out;
    box-shadow: 0 4px 15px rgba(16, 172, 132, 0.3);
}

.error-item {
    background: linear-gradient(135deg, #ff6b6b 0%, #ee5a24 100%);
    color: white;
    padding: 1rem 1.5rem;
    border-radius: 10px;
    margin: 0.5rem 0;
    display: flex;
    align-items: center;
    gap: 1rem;
    animation: slideInRight 0.5s ease-out;
    box-shadow: 0 4px 15px rgba(255, 107, 107, 0.3);
}

.warning-item {
    background: linear-gradient(135deg, #feca57 0%, #ff9f43 100%);
    color: white;
    padding: 1rem 1.5rem;
    border-radius: 10px;
    margin: 0.5rem 0;
    display: flex;
    align-items: center;
    gap: 1rem;
    animation: slideInLeft 0.5s ease-out;
    box-shadow: 0 4px 15px rgba(254, 202, 87, 0.3);
}

/* Upload section */
.upload-section {
    background: linear-gradient(135deg, #74b9ff 0%, #0984e3 100%);
    padding: 2rem;
    border-radius: 15px;
    text-align: center;
    color: white;
    margin: 2rem 0;
    box-shadow: 0 8px 32px rgba(116, 185, 255, 0.3);
}

/* Progress container */
.progress-container {
    background: rgba(255,255,255,0.9);
    border-radius: 15px;
    padding: 1.5rem;
    margin: 2rem 0;
    box-shadow: 0 8px 32px rgba(0,0,0,0.1);
}

/* Animations */
@keyframes fadeInUp {
    from {
        opacity: 0;
        transform: translateY(30px);
    }
    to {
        opacity: 1;
        transform: translateY(0);
    }
}

@keyframes slideInLeft {
    from {
        opacity: 0;
        transform: translateX(-30px);
    }
    to {
        opacity: 1;
        transform: translateX(0);
    }
}

@keyframes slideInRight {
    from {
        opacity: 0;
        transform: translateX(30px);
    }
    to {
        opacity: 1;
        transform: translateX(0);
    }
}

/* Sidebar styles */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #2d3748 0%, #1a202c 100%);
}

[data-testid="stSidebar"] * {
    color: white !important;
}

/* Enhanced button styles */
.stDownloadButton > button {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.75rem 1.5rem !important;
    font-weight: 600 !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3) !important;
    width: 100% !important;
    margin: 0.25rem 0 !important;
}

.stDownloadButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(102, 126, 234, 0.4) !important;
}

/* Tab styling */
.stTabs [data-baseweb="tab-list"] {
    gap: 1rem;
    background: rgba(255,255,255,0.1);
    border-radius: 15px;
    padding: 0.5rem;
}

.stTabs [data-baseweb="tab"] {
    border-radius: 10px;
    padding: 1rem 1.5rem;
    font-weight: 600;
    color: white;
}

/* Metrics styling */
[data-testid="metric-container"] {
    background: rgba(255,255,255,0.9);
    border-radius: 15px;
    padding: 1rem;
    box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    text-align: center;
}

/* File uploader styling */
[data-testid="stFileUploader"] {
    background: rgba(255,255,255,0.9);
    border-radius: 15px;
    padding: 2rem;
    border: 2px dashed #667eea;
}

/* Status text styling */
.status-text {
    text-align: center;
    padding: 1rem;
    background: rgba(255,255,255,0.9);
    border-radius: 10px;
    margin: 1rem 0;
    font-weight: 600;
    color: #2d3748;
}

/* Responsive design */
@media (max-width: 768px) {
    .hero-title {
        font-size: 2.5rem;
    }
    
    .hero-stats {
        gap: 1.5rem;
    }
    
    .stat-number {
        font-size: 2rem;
    }
    
    .custom-card {
        padding: 1.5rem;
    }
}
</style>
""", unsafe_allow_html=True)

# --- Configuration / Checklist ---
CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association", 
        "Board Resolution",
        "UBO Declaration Form",
        "Register of Members and Directors"
    ]
}

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa", "articles"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Board Resolution": ["board resolution", "resolution of the board", "directors resolution"],
    "UBO Declaration Form": ["ubo", "ultimate beneficial owner", "ubo declaration", "beneficial owner"],
    "Register of Members and Directors": ["register of members", "register of directors", "register of members and directors"],
    "Incorporation Application Form": ["incorporation application", "application for incorporation"],
    "Shareholder Resolution Templates": ["shareholder resolution", "shareholder resolution template"]
}

# --- Utility functions ---
@st.cache_data
def docx_to_text_bytesio(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes with error handling"""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        texts = []
        
        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text.strip())
        
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)
        
        return "\n".join(texts)
    except Exception as e:
        st.error(f"Error reading document: {str(e)}")
        return ""

def detect_doc_type(text: str, filename: str = "") -> str:
    """Enhanced document type detection"""
    if not text.strip():
        return "Unknown"
        
    lower_text = text.lower()
    lower_filename = filename.lower()
    
    # Score based on content and filename
    scores = {}
    for dtype, keywords in DOC_TYPE_KEYWORDS.items():
        content_score = sum(3 if kw in lower_text else 0 for kw in keywords)
        filename_score = sum(2 if kw in lower_filename else 0 for kw in keywords)
        total_score = content_score + filename_score
        
        if total_score > 0:
            scores[dtype] = total_score
    
    if not scores:
        # Fallback detection
        if any(term in lower_text for term in ["articles", "association"]):
            return "Articles of Association"
        elif any(term in lower_text for term in ["memorandum", "moa"]):
            return "Memorandum of Association"
        elif any(term in lower_text for term in ["resolution", "resolved"]):
            return "Board Resolution"
        return "Unknown"
    
    return max(scores.items(), key=lambda x: x[1])[0]

def check_jurisdiction(text: str) -> List[str]:
    """Check for jurisdiction-related issues"""
    flags = []
    lower = text.lower()
    
    # Check for incorrect jurisdiction references
    federal_patterns = [
        "uae federal court", "federal courts", "uae federal", 
        "dubai courts", "abu dhabi courts", "sharjah courts"
    ]
    
    for pattern in federal_patterns:
        if pattern in lower and "adgm" not in lower:
            flags.append(f"Document references '{pattern}' instead of ADGM jurisdiction.")
    
    return flags

def check_missing_signatory(text: str) -> List[str]:
    """Check for missing signature blocks"""
    flags = []
    lower = text.lower()
    
    signature_indicators = [
        "signature", "signed", "for and on behalf", 
        "director", "secretary", "authorized signatory"
    ]
    
    if not any(indicator in lower for indicator in signature_indicators):
        flags.append("No signature block or signatory section detected.")
    
    return flags

def check_ambiguous_language(text: str) -> List[str]:
    """Check for ambiguous or weak language"""
    flags = []
    ambiguous_terms = [
        "may", "could", "might", "endeavor", "endeavour", 
        "best efforts", "reasonable endeavours", "attempt to"
    ]
    
    found_terms = []
    for term in ambiguous_terms:
        if re.search(r"\b" + re.escape(term) + r"\b", text.lower()):
            found_terms.append(term)
    
    if found_terms:
        flags.append(f"Ambiguous/optional language detected: {', '.join(sorted(set(found_terms)))}")
    
    return flags

def create_annotated_docx(original_bytes: bytes, flags_report: List[Dict]) -> bytes:
    """Create annotated DOCX with review comments"""
    try:
        doc = Document(io.BytesIO(original_bytes))
        doc.add_page_break()
        
        # Add review section
        try:
            heading = doc.add_heading("ADGM Compliance Review Comments", level=1)
            heading.style.font.color.rgb = None  # Keep default color
        except Exception:
            # Fallback for heading creation
            paragraph = doc.add_paragraph("ADGM Compliance Review Comments")
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("ADGM Compliance Review Comments")
            run.bold = True
            try:
                run.font.size = Pt(16)
            except Exception:
                pass
        
        # Add timestamp
        timestamp_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
        try:
            if hasattr(timestamp_para.style, 'font') and hasattr(timestamp_para.style.font, 'size'):
                timestamp_para.style.font.size = Pt(10)
        except Exception:
            pass
        
        doc.add_paragraph()  # Empty line
        
        if not flags_report:
            doc.add_paragraph("✅ No compliance issues found. Document appears to be compliant with ADGM requirements.")
        else:
            # Group issues by severity
            severity_order = {'High': 1, 'Medium': 2, 'Low': 3}
            sorted_issues = sorted(flags_report, key=lambda x: severity_order.get(x.get('severity', 'Low'), 3))
            
            for i, issue in enumerate(sorted_issues, start=1):
                # Issue header
                issue_para = doc.add_paragraph()
                issue_run = issue_para.add_run(f"{i}. {issue.get('severity', 'Unknown')} Priority Issue")
                issue_run.bold = True
                
                # Document name
                doc.add_paragraph(f"Document: {issue.get('document', 'Unknown')}")
                
                # Location
                location = issue.get('location_hint', 'Not specified')
                doc.add_paragraph(f"Location: {location}")
                
                # Issue description
                issue_desc = issue.get('issue', 'No description provided')
                doc.add_paragraph(f"Issue: {issue_desc}")
                
                # Suggestion
                suggestion = issue.get('suggestion', 'No suggestion provided')
                doc.add_paragraph(f"Recommendation: {suggestion}")
                
                doc.add_paragraph()  # Empty line between issues
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
        
    except Exception as e:
        st.error(f"Error creating annotated document: {str(e)}")
        return original_bytes

def analyze_document_text(docname: str, text: str) -> List[Dict]:
    """Comprehensive document analysis"""
    if not text.strip():
        return [{
            "document": docname,
            "paragraph_index": 0,
            "location_hint": "Entire document",
            "issue": "Document appears to be empty or unreadable",
            "severity": "High",
            "suggestion": "Please ensure the document contains text and is not corrupted"
        }]
    
    issues = []
    paragraphs = [p.strip() for p in text.splitlines() if p.strip()]

    # Jurisdiction checks
    jurisdiction_flags = check_jurisdiction(text)
    for flag in jurisdiction_flags:
        # Find specific paragraph with jurisdiction issue
        paragraph_index = 0
        for idx, para in enumerate(paragraphs):
            if "federal" in para.lower() or "court" in para.lower():
                paragraph_index = idx
                break
        
        issues.append({
            "document": docname,
            "paragraph_index": paragraph_index,
            "location_hint": f"Paragraph {paragraph_index + 1}" if paragraph_index >= 0 else "Jurisdiction clause",
            "issue": flag,
            "severity": "High",
            "suggestion": "Replace reference to UAE Federal Courts with ADGM Courts and specify ADGM jurisdiction."
        })

    # Signatory checks
    signatory_flags = check_missing_signatory(text)
    for flag in signatory_flags:
        issues.append({
            "document": docname,
            "paragraph_index": len(paragraphs) - 1 if paragraphs else 0,
            "location_hint": "End of document",
            "issue": flag,
            "severity": "Medium",
            "suggestion": "Add a clear signature block with 'For and on behalf of [Company Name]' and designated signatory fields."
        })

    # Ambiguous language checks
    ambiguous_flags = check_ambiguous_language(text)
    for flag in ambiguous_flags:
        issues.append({
            "document": docname,
            "paragraph_index": 0,
            "location_hint": "Throughout document",
            "issue": flag,
            "severity": "Low",
            "suggestion": "Replace ambiguous terms with clear, definitive obligations (e.g., 'shall' instead of 'may')."
        })

    # ADGM jurisdiction clause check
    if not re.search(r"\b(adgm|abu dhabi global market)\b", text.lower()):
        issues.append({
            "document": docname,
            "paragraph_index": 0,
            "location_hint": "Jurisdiction clause (missing)",
            "issue": "No explicit ADGM jurisdiction clause detected.",
            "severity": "High",
            "suggestion": "Add a jurisdiction clause stating 'This document shall be governed by the laws of ADGM and disputes shall be resolved by ADGM Courts.'"
        })

    return issues

# --- Hero Section ---
st.markdown("""
<div class="hero-section">
    <div class="hero-title">⚖️ ADGM Corporate Agent</div>
    <div class="hero-subtitle">
        AI-Powered Legal Document Review & Compliance Validation
    </div>
    <div class="hero-stats">
        <div class="stat-item">
            <span class="stat-number">100%</span>
            <span class="stat-label">Automated</span>
        </div>
        <div class="stat-item">
            <span class="stat-number">24/7</span>
            <span class="stat-label">Available</span>
        </div>
        <div class="stat-item">
            <span class="stat-number">⚡</span>
            <span class="stat-label">Instant</span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# --- Sidebar ---
with st.sidebar:
    st.markdown("### 🎯 **Features**")
    st.markdown("""
    ✨ **Smart Document Detection**  
    🔍 **Compliance Validation**  
    ⚠️ **Risk Assessment**  
    📋 **Automated Checklists**  
    📄 **Annotated Reports**  
    💾 **Instant Downloads**  
    """)
    
    st.markdown("---")
    st.markdown("### 📊 **Required Documents**")
    for doc_type in CHECKLISTS["Company Incorporation"]:
        st.markdown(f"• {doc_type}")
    
    st.markdown("---")
    st.markdown("### 💡 **Tips**")
    st.markdown("""
    📝 Upload .docx files only  
    🔍 Review all flagged issues  
    ⚖️ Ensure ADGM jurisdiction  
    ✍️ Check signature blocks  
    """)

# --- Main Content ---
st.markdown('<div class="upload-section">', unsafe_allow_html=True)
st.markdown("### 📤 **Upload Your Documents**")
st.markdown("*Drag & drop your .docx files below for instant ADGM compliance review*")
st.markdown('</div>', unsafe_allow_html=True)

uploaded = st.file_uploader(
    "Choose DOCX files",
    accept_multiple_files=True,
    type=["docx"],
    help="Upload one or more .docx files for comprehensive ADGM compliance analysis"
)

if uploaded:
    st.markdown('<div class="progress-container">', unsafe_allow_html=True)
    
    # Progress tracking
    progress_bar = st.progress(0)
    status_placeholder = st.empty()
    
    with st.spinner("🔮 AI is analyzing your documents..."):
        processed_docs = []
        documents_present = []
        all_issues = []
        
        total_files = len(uploaded)
        
        for i, uploaded_file in enumerate(uploaded):
            # Update progress
            progress = int(((i + 1) / total_files) * 100)
            progress_bar.progress(progress / 100.0)
            
            with status_placeholder:
                st.markdown(f'<div class="status-text">📄 Processing {uploaded_file.name}... ({i+1}/{total_files})</div>', unsafe_allow_html=True)
            
            try:
                # Reset file pointer and read content
                uploaded_file.seek(0)
                raw_bytes = uploaded_file.read()
                
                # Extract text
                text_content = docx_to_text_bytesio(raw_bytes)
                
                # Detect document type
                doc_type = detect_doc_type(text_content, uploaded_file.name)
                
                # Add to present documents
                if doc_type != "Unknown":
                    documents_present.append(doc_type)
                else:
                    # Use filename as fallback
                    base_name = os.path.splitext(uploaded_file.name)[0]
                    documents_present.append(base_name)
                
                # Analyze document
                document_issues = analyze_document_text(uploaded_file.name, text_content)
                all_issues.extend(document_issues)
                
                # Store processed document info
                processed_docs.append({
                    "filename": uploaded_file.name,
                    "type_detected": doc_type,
                    "issues": document_issues,
                    "raw_bytes": raw_bytes,
                    "text_content": text_content
                })
                
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                # Add error as an issue
                all_issues.append({
                    "document": uploaded_file.name,
                    "paragraph_index": 0,
                    "location_hint": "File processing",
                    "issue": f"Failed to process document: {str(e)}",
                    "severity": "High",
                    "suggestion": "Please ensure the file is a valid .docx document and try again."
                })
            
            # Small delay for visual effect
            time.sleep(0.2)
        
        # Complete progress
        progress_bar.progress(1.0)
        with status_placeholder:
            st.markdown('<div class="status-text">✅ Analysis complete!</div>', unsafe_allow_html=True)
        
        time.sleep(0.6)
        progress_bar.empty()
        status_placeholder.empty()
    
    st.markdown('</div>', unsafe_allow_html=True)

    # Calculate completion metrics
    required_docs = CHECKLISTS["Company Incorporation"]
    present_types = set(documents_present)
    missing_docs = [doc for doc in required_docs if doc not in present_types]
    
    # Generate comprehensive report
    report = {
        "analysis_metadata": {
            "process_type": "Company Incorporation",
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "total_documents_analyzed": len(uploaded),
            "analysis_version": "1.0"
        },
        "document_summary": {
            "documents_uploaded": len(uploaded),
            "required_documents": len(required_docs),
            "uploaded_document_types": list(present_types),
            "missing_documents": missing_docs,
            "completion_percentage": round(((len(required_docs) - len(missing_docs)) / len(required_docs)) * 100, 1)
        },
        "compliance_analysis": {
            "total_issues": len(all_issues),
            "high_severity_issues": len([i for i in all_issues if i.get('severity') == 'High']),
            "medium_severity_issues": len([i for i in all_issues if i.get('severity') == 'Medium']),
            "low_severity_issues": len([i for i in all_issues if i.get('severity') == 'Low']),
            "issues_by_document": {doc["filename"]: len(doc["issues"]) for doc in processed_docs}
        },
        "detailed_issues": all_issues
    }

    # Results Summary Section
    st.markdown("## 📊 **Analysis Summary**")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "📄 Documents Analyzed", 
            len(uploaded),
            help="Total number of documents processed"
        )
    
    with col2:
        high_issues = report["compliance_analysis"]["high_severity_issues"]
        delta_color = "inverse" if high_issues > 0 else "normal"
        st.metric(
            "🚨 Critical Issues", 
            high_issues,
            delta=f"{high_issues} found" if high_issues > 0 else "None found",
            delta_color=delta_color,
            help="High priority compliance issues that require immediate attention"
        )
    
    with col3:
        completion_rate = report["document_summary"]["completion_percentage"]
        st.metric(
            "✅ Document Completeness", 
            f"{completion_rate}%",
            delta=f"{len(required_docs) - len(missing_docs)}/{len(required_docs)} docs",
            help="Percentage of required documents provided"
        )
    
    with col4:
        total_issues = len(all_issues)
        st.metric(
            "⚠️ Total Issues", 
            total_issues,
            help="All compliance issues found across all documents"
        )

    # Enhanced Tabs with better organization
    tab1, tab2, tab3, tab4 = st.tabs([
        "📋 **Document Checklist**", 
        "⚠️ **Issues & Recommendations**", 
        "📥 **Reports & Downloads**", 
        "📊 **Detailed Analysis**"
    ])

    # ---------------------------
    # Tab 1: Document Checklist
    # ---------------------------
    with tab1:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("### 📋 **ADGM Compliance Checklist**")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            for doc_type in required_docs:
                if doc_type in present_types:
                    st.markdown(f'<div class="success-item">✅ <strong>{doc_type}</strong> - Document Found & Analyzed</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="error-item">❌ <strong>{doc_type}</strong> - Missing Document</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown("**📈 Completion Status**")
            completion_percentage = report["document_summary"]["completion_percentage"]
            
            if completion_percentage == 100:
                st.success("🎉 All required documents provided!")
            elif completion_percentage >= 80:
                st.warning(f"⚠️ {completion_percentage}% complete - Almost there!")
            else:
                st.error(f"🚨 {completion_percentage}% complete - More documents needed")
        
        if missing_docs:
            st.markdown("### 📝 **Action Required**")
            st.error(f"Please upload the following missing documents: **{', '.join(missing_docs)}**")
        
        st.markdown('</div>', unsafe_allow_html=True)

    # --------------------------------
    # Tab 2: Issues & Recommendations
    # --------------------------------
    with tab2:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("### ⚠️ **Compliance Issues & Recommendations**")
        
        if not all_issues:
            st.success("🎉 Excellent! No compliance issues detected. Your documents appear to meet ADGM requirements.")
        else:
            # Group issues by severity with enhanced display
            high_issues = [i for i in all_issues if i.get('severity') == 'High']
            medium_issues = [i for i in all_issues if i.get('severity') == 'Medium']
            low_issues = [i for i in all_issues if i.get('severity') == 'Low']
            
            # Display issue summary
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"**🚨 High Priority:** {len(high_issues)}")
            with col2:
                st.markdown(f"**⚠️ Medium Priority:** {len(medium_issues)}")
            with col3:
                st.markdown(f"**ℹ️ Low Priority:** {len(low_issues)}")
            
            st.markdown("---")
            
            if high_issues:
                st.markdown("#### 🚨 **Critical Issues - Immediate Action Required**")
                for idx, issue in enumerate(high_issues, 1):
                    st.markdown(f'''
                    <div class="error-item">
                        <div>
                            <strong>🚨 Issue #{idx}: {issue["document"]}</strong><br>
                            <strong>Location:</strong> {issue["location_hint"]}<br>
                            <strong>Problem:</strong> {issue["issue"]}<br>
                            <strong>💡 Recommendation:</strong> {issue["suggestion"]}
                        </div>
                    </div>
                    ''', unsafe_allow_html=True)
            
            if medium_issues:
                st.markdown("#### ⚠️ **Medium Priority Issues - Should Address Soon**")
                for idx, issue in enumerate(medium_issues, 1):
                    st.markdown(f'''
                    <div class="warning-item">
                        <div>
                            <strong>⚠️ Issue #{idx}: {issue["document"]}</strong><br>
                            <strong>Location:</strong> {issue["location_hint"]}<br>
                            <strong>Problem:</strong> {issue["issue"]}<br>
                            <strong>💡 Recommendation:</strong> {issue["suggestion"]}
                        </div>
                    </div>
                    ''', unsafe_allow_html=True)
            
            if low_issues:
                st.markdown("#### ℹ️ **Minor Issues - Consider Addressing**")
                with st.expander(f"Show {len(low_issues)} minor issues"):
                    for idx, issue in enumerate(low_issues, 1):
                        st.markdown(f'''
                        <div class="success-item">
                            <div>
                                <strong>ℹ️ Issue #{idx}: {issue["document"]}</strong><br>
                                <strong>Location:</strong> {issue["location_hint"]}<br>
                                <strong>Problem:</strong> {issue["issue"]}<br>
                                <strong>💡 Recommendation:</strong> {issue["suggestion"]}
                            </div>
                        </div>
                        ''', unsafe_allow_html=True)

            # ---------------------------
            # Tab2: Issues Table View
            # ---------------------------
            st.markdown("### 🔎 Issues Table")
            # Build DataFrame for better scanning/exporting
            issues_df = pd.DataFrame([
                {
                    'Document': issue.get('document'),
                    'Severity': issue.get('severity'),
                    'SeverityBadge': ("🚨 High" if issue.get('severity') == 'High'
                                      else "⚠️ Medium" if issue.get('severity') == 'Medium'
                                      else "ℹ️ Low"),
                    'Location': issue.get('location_hint'),
                    'Issue': issue.get('issue'),
                    'Recommendation': issue.get('suggestion')
                }
                for issue in all_issues
            ])
            if not issues_df.empty:
                # show the table (streamlit will render it interactively)
                st.dataframe(issues_df[['Document', 'SeverityBadge', 'Location', 'Issue', 'Recommendation']], height=300)
                
                # CSV download
                st.download_button(
                    "📥 Download Issues CSV",
                    data=issues_df.to_csv(index=False).encode('utf-8'),
                    file_name=f"adgm_issues_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
        
        st.markdown('</div>', unsafe_allow_html=True)

    # ---------------------------
    # Tab 3: Reports & Downloads
    # ---------------------------
    with tab3:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("### 📥 **Reports & Download Center**")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📊 **Compliance Reports**")
            
            # JSON Report
            json_report = json.dumps(report, indent=2, ensure_ascii=False)
            st.download_button(
                "📄 Download Complete JSON Report",
                data=json_report.encode('utf-8'),
                file_name=f"adgm_compliance_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                help="Comprehensive analysis report in JSON format"
            )
            
            # CSV Summary Report
            if all_issues:
                # Create issues DataFrame (reuse)
                issues_df = pd.DataFrame([
                    {
                        'Document': issue['document'],
                        'Severity': issue['severity'],
                        'Location': issue['location_hint'],
                        'Issue': issue['issue'],
                        'Recommendation': issue['suggestion']
                    }
                    for issue in all_issues
                ])
                
                csv_data = issues_df.to_csv(index=False)
                st.download_button(
                    "📊 Download Issues Summary (CSV)",
                    data=csv_data,
                    file_name=f"adgm_issues_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    help="Issues summary in spreadsheet format"
                )
        
        with col2:
            st.markdown("#### 📝 **Annotated Documents**")
            
            if processed_docs:
                for doc in processed_docs:
                    try:
                        annotated_bytes = create_annotated_docx(doc["raw_bytes"], doc["issues"])
                        
                        # Create safe filename
                        base_name = os.path.splitext(doc['filename'])[0]
                        safe_filename = f"{base_name}_ADGM_Review.docx"
                        
                        issues_count = len(doc["issues"])
                        button_text = f"📎 {doc['filename']}"
                        if issues_count > 0:
                            button_text += f" ({issues_count} issues)"
                        
                        st.download_button(
                            button_text,
                            data=annotated_bytes,
                            file_name=safe_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            help=f"Original document with compliance review comments added"
                        )
                    except Exception as e:
                        st.error(f"Error preparing {doc['filename']}: {str(e)}")
            else:
                st.info("No documents available for download")
        
        # Quick Actions Section
        st.markdown("---")
        st.markdown("#### ⚡ **Quick Actions**")
        
        action_col1, action_col2, action_col3 = st.columns(3)
        
        with action_col1:
            if st.button("🔄 **Re-analyze Documents**", help="Run analysis again with current documents"):
                st.experimental_rerun()
        
        with action_col2:
            if st.button("📋 **Print Checklist**", help="Generate printable checklist"):
                checklist_content = "ADGM Company Incorporation Checklist\n\n"
                for i, doc in enumerate(required_docs, 1):
                    status = "✅ Complete" if doc in present_types else "❌ Missing"
                    checklist_content += f"{i}. {doc}: {status}\n"
                
                st.download_button(
                    "📄 Download Printable Checklist",
                    data=checklist_content,
                    file_name=f"ADGM_Checklist_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain"
                )
        
        with action_col3:
            if len(missing_docs) > 0:
                st.info(f"📝 Still need: {len(missing_docs)} documents")
            else:
                st.success("✅ All documents provided!")
        
        st.markdown('</div>', unsafe_allow_html=True)

    # ---------------------------
    # Tab 4: Detailed Analysis
    # ---------------------------
    with tab4:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("### 📊 **Detailed Analysis Report**")
        
        # Enhanced metrics dashboard
        st.markdown("#### 📈 **Analysis Metrics**")
        
        metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)
        
        with metric_col1:
            st.metric(
                "Documents Processed", 
                len(processed_docs),
                help="Total number of documents successfully analyzed"
            )
        
        with metric_col2:
            st.metric(
                "Issues Identified", 
                len(all_issues),
                help="Total compliance issues found across all documents"
            )
        
        with metric_col3:
            st.metric(
                "Document Types Detected", 
                len(set(doc['type_detected'] for doc in processed_docs if doc['type_detected'] != 'Unknown')),
                help="Number of different document types successfully identified"
            )
        
        with metric_col4:
            avg_issues = round(len(all_issues) / len(processed_docs), 1) if processed_docs else 0
            st.metric(
                "Avg Issues per Doc", 
                avg_issues,
                help="Average number of issues per document"
            )
        
        # Document-by-document breakdown
        st.markdown("#### 📋 **Document Analysis Breakdown**")
        
        # Build a concise summary table
        summary_rows = []
        for doc in processed_docs:
            summary_rows.append({
                "Filename": doc["filename"],
                "Type Detected": doc["type_detected"],
                "Issues Found": len(doc["issues"]),
                "Text Length": len(doc.get("text_content", "")),
            })
        summary_df = pd.DataFrame(summary_rows)
        if not summary_df.empty:
            st.dataframe(summary_df, height=220)
        
        for doc in processed_docs:
            with st.expander(f"📄 {doc['filename']} - {doc['type_detected']} ({len(doc['issues'])} issues)"):
                col1, col2 = st.columns([1, 2])
                
                with col1:
                    st.markdown("**Document Info:**")
                    st.markdown(f"• **Type:** {doc['type_detected']}")
                    st.markdown(f"• **Issues Found:** {len(doc['issues'])}")
                    st.markdown(f"• **Text Length:** {len(doc.get('text_content', ''))} characters")
                
                with col2:
                    if doc['issues']:
                        st.markdown("**Issues in this document:**")
                        for i, issue in enumerate(doc['issues'], 1):
                            severity_emoji = {"High": "🚨", "Medium": "⚠️", "Low": "ℹ️"}.get(issue['severity'], "•")
                            st.markdown(f"{severity_emoji} **{issue['severity']}:** {issue['issue']}")
                    else:
                        st.success("✅ No issues found in this document!")
        
        # Raw JSON Report Section
        st.markdown("#### 🔍 **Complete Technical Report**")
        with st.expander("📄 **View Full JSON Report**", expanded=False):
            st.json(report)
        
        # Export options
        st.markdown("#### 💾 **Advanced Export Options**")
        
        export_col1, export_col2 = st.columns(2)
        
        with export_col1:
            # Summary report
            summary_report = {
                "summary": report["document_summary"],
                "compliance": report["compliance_analysis"],
                "timestamp": report["analysis_metadata"]["generated_at"]
            }
            
            st.download_button(
                "📊 Download Executive Summary",
                data=json.dumps(summary_report, indent=2),
                file_name=f"adgm_executive_summary_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                help="High-level summary report for management"
            )
        
        with export_col2:
            # Technical report
            technical_report = {
                "metadata": report["analysis_metadata"],
                "detailed_issues": report["detailed_issues"],
                "document_analysis": [
                    {
                        "filename": doc["filename"],
                        "type_detected": doc["type_detected"],
                        "issues_count": len(doc["issues"]),
                        "issues": doc["issues"]
                    }
                    for doc in processed_docs
                ]
            }
            
            st.download_button(
                "🔧 Download Technical Report",
                data=json.dumps(technical_report, indent=2),
                file_name=f"adgm_technical_report_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                help="Detailed technical analysis for legal review"
            )
        
        st.markdown('</div>', unsafe_allow_html=True)

else:
    # Enhanced Welcome Section
    st.markdown('<div class="custom-card">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 🚀 **Welcome to ADGM Corporate Agent**")
        st.markdown("""
        Our AI-powered platform helps you ensure your corporate documents comply with 
        **Abu Dhabi Global Market (ADGM)** regulations. Simply upload your .docx files 
        and receive instant compliance analysis with actionable recommendations.
        """)
        
        st.markdown("### 📋 **How It Works**")
        
        step_col1, step_col2, step_col3 = st.columns(3)
        
        with step_col1:
            st.markdown("""
            **1. Upload** 📤  
            • Drag & drop .docx files
            • Multiple files supported
            • Secure processing
            """)
        
        with step_col2:
            st.markdown("""
            **2. AI Analysis** 🔍  
            • Document type detection
            • Compliance validation
            • Risk assessment
            """)
        
        with step_col3:
            st.markdown("""
            **3. Get Results** 📥  
            • Detailed reports
            • Annotated documents
            • Action recommendations
            """)
    
    with col2:
        st.markdown("### 📊 **What We Check**")
        st.markdown("""
        ✅ **Jurisdiction Clauses**  
        ✅ **ADGM Compliance**  
        ✅ **Signature Blocks**  
        ✅ **Legal Language**  
        ✅ **Document Completeness**  
        ✅ **Required Forms**  
        """)
        
        st.markdown("---")
        st.info("👆 **Ready to start?** Upload your documents above!")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Feature Highlights
    st.markdown("### ✨ **Key Features**")
    
    feature_col1, feature_col2, feature_col3 = st.columns(3)
    
    with feature_col1:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("#### 🤖 **Smart AI Detection**")
        st.markdown("""
        Our AI automatically identifies document types and detects 
        compliance issues specific to ADGM requirements.
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with feature_col2:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("#### ⚡ **Instant Results**")
        st.markdown("""
        Get comprehensive compliance reports in seconds, not hours. 
        Perfect for busy legal teams and corporate professionals.
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with feature_col3:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("#### 📄 **Professional Reports**")
        st.markdown("""
        Download annotated documents and detailed reports ready 
        for legal review and regulatory submission.
        """)
        st.markdown('</div>', unsafe_allow_html=True)

# --- Enhanced Footer ---
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 2rem; background: rgba(255,255,255,0.1); border-radius: 15px; margin-top: 2rem;">
    <div style="color: #2d3748; font-size: 1.1rem; font-weight: 600; margin-bottom: 1rem;">
        🏛️ <strong>ADGM Corporate Agent</strong> | Powered by Advanced AI
    </div>
    <div style="color: #4a5568; font-size: 0.9rem;">
        ⚡ Lightning-fast Analysis • 🔒 Secure Processing • 📊 Professional Reports • 🎯 ADGM Focused
    </div>
    <div style="color: #718096; font-size: 0.8rem; margin-top: 1rem;">
        Built with Streamlit • Enhanced for Legal Professionals • Last Updated: {timestamp}
    </div>
</div>
""".format(timestamp=datetime.now().strftime('%B %Y')), unsafe_allow_html=True)
